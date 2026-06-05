"""生成 游戏对象 / UI界面 / 驱动工具 的 Word 文档"""
import os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn

OUT = "docs"
os.makedirs(OUT, exist_ok=True)

def new_doc():
    d = Document()
    s = d.styles['Normal']
    s.font.name = "Consolas"; s.font.size = Pt(9)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return d

def h0(d,t): d.add_heading(t, level=0)
def h1(d,t): d.add_heading(t, level=1)
def h2(d,t): d.add_heading(t, level=2)
def h3(d,t): d.add_heading(t, level=3)
def para(d,t):
    p=d.add_paragraph(t); p.paragraph_format.space_after=Pt(3)
def code(d,t):
    for line in t.split("\n"):
        p=d.add_paragraph(); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0); p.paragraph_format.left_indent=Cm(0.4)
        r=p.add_run(line); r.font.size=Pt(7); r.font.name="Consolas"

def add_section(d, fname, content):
    d.add_heading(fname, level=1)
    parts = content.split("\n## ")
    for idx, part in enumerate(parts):
        part = part.strip()
        if not part: continue
        if not part.startswith("##") and ("概述" in part.lower() or idx>0):
            h2(d, "概述"); para(d, part.replace("## 概述","").strip()); continue
        lines = part.split("\n"); i = 0; in_code = False; code_buf = []
        while i < len(lines):
            line = lines[i]
            if line.startswith("### "):
                if in_code: code(d,"\n".join(code_buf)); code_buf=[]; in_code=False
                h3(d, line[4:].strip()); i+=1; continue
            if line.strip() == "```c":
                if in_code: code(d,"\n".join(code_buf)); code_buf=[]; in_code=False
                else: in_code=True
                i+=1; continue
            if in_code: code_buf.append(line)
            else:
                s=line.strip()
                if s and not s.startswith("#"):
                    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(2)
                    for pb in re.split(r'(\*\*.+?\*\*)', s):
                        if pb.startswith("**") and pb.endswith("**"):
                            r=p.add_run(pb[2:-2]); r.bold=True
                        else: p.add_run(pb)
            i+=1
        if in_code and code_buf: code(d,"\n".join(code_buf))

# ========================================================
#  文档3: 游戏对象
# ========================================================
def make_obj_docs():
    d = new_doc()
    h0(d, "Beat Plane — 游戏对象代码详解")

    files = {
"obj/player/player.c": """## 概述
玩家对象核心实现。player_t继承game_obj_t，包含4种可切换飞机的完整生命周期：外观切换、HP管理、移动输入、子弹射击、X/Y双技能系统、护盾/加速状态、技能CD可视化时间戳。

## 关键数据结构
### player_t
```c
typedef struct {
    game_obj_t base;       // 继承基类(x,y,vx,vy,active,type,apr...)
    int16_t hp, hp_max;
    lv_obj_t *hp_bar;     // HP条LVGL对象
    int current_plane_id;
    int16_t bullet_damage, bullet_vx, bullet_vy;
    apr_id_t bullet_apr;
    uint16_t skill_x_cd, skill_y_cd;
    void (*skill_x_active)(void);  // 函数指针，指向当前飞机的X技能
    void (*skill_y_active)(void);
    lv_obj_t *shield_overlay; bool shield_active;
    bool speed_boost_active;
    uint32_t skill_x_last_use, skill_y_last_use; // CD可视化用
} player_t;
```

### plane_config_t (静态配置表)
4架飞机定义：Player(HP200,射速200ms,伤害34)、Ember(HP200,射速150ms,伤害80)、Stream(HP200,射速250ms,伤害10)、Verdant(HP250,射速220ms,伤害15)

### player_init
- **逻辑流程**: 分配player_t内存→设置默认位置(512,500)→从plane_configs[0]拷贝属性→创建HP条和LVGL图像→创建护盾半透明遮罩→注册按键(A=射击,X=E键技能,Y=F键技能)和4个事件

### player_apply_config
切换飞机时：取消旧按键回调→重置护盾/加速/CD时间戳→从配置表拷贝新飞机的全部属性→重新注册按键(含CD)

### player_update (每帧)
状态检查→Verdant检测X键长按开启加速→摇杆读取(-7~7)→加速时速度翻倍→player_move

### player_move
坐标+=vx/vy→边界钳制→lv_obj_set_pos更新位置→护盾遮罩跟随

## 技能系统
### X键技能 (按E触发)
- **player_skill_triple_shot**: 发射3颗子弹(中间+左右偏3)
- **player_skill_burn_bullet**: 发射带BULLET_FLAG_BURN标志的子弹(Ember)
- **player_skill_freeze_bullet**: 发射带BULLET_FLAG_FREEZE标志的子弹(Stream)

### Y键技能 (按F触发)
- **player_skill_shield**: 1秒护盾无敌(Player)
- **player_skill_flame_wall**: 发射火墙(Ember)
- **player_skill_bullet_slow**: 全屏敌方子弹减速2秒+冰冻蓝色渲染(Stream)
- **player_skill_hp_reclaim**: 恢复50HP(Verdant)
""",

"obj/bullet/bullet.c": """## 概述
子弹系统实现。bullet_t继承game_obj_t，基于对象池(MAX_BULLET_COUNT=50)管理。支持普通子弹和特殊效果子弹。

### bullet_init
清零数组→初始化对象池→预创建50个LVGL图像对象→挂载update/show/hide函数→注册EVENT_BULLET_HIT_ENEMY和EVENT_BULLET_HIT_PLAYER事件

### bullet_create
从池分配→设置source/damage/vx/vy/x/y/flags=0→apr_apply换肤→重置透明度(LV_OPA_COVER)→显示→若减速激活且来源是敌人则应用半透明冰冻效果
```c
if (enemy_slow_active && source != NULL && source->type == GAME_OBJ_TYPE_ENEMY) {
    lv_obj_set_style_opa(bullets[index].base.obj, LV_OPA_60, 0);
}
```

### bullet_move
计算有效速度→若减速激活且来源是敌人则速度减半→更新坐标→出界销毁

### bullet_hide
隐藏图像→旋转归零(防追踪弹角度残留)→归还对象池

### 标志位系统
BULLET_FLAG_NONE(0x00) | BURN(0x01灼烧) | FREEZE(0x02冻结) | REFLECTED(0x04被护盾反射)

### bullet_set_enemy_slow
设置全局减速开关→批量遍历50发子弹→对敌方子弹应用/移除冰冻半透明效果
""",

"obj/bullet/bullet_behaviors.c": """## 概述
3种子弹行为算法：圆周运动(泰勒近似sin/cos)、正弦波动、追踪玩家。

### bullet_behave_circle
用theta=0.12作固定角速度，泰勒展开前两项近似sin/cos，旋转矩阵更新速度方向，归一化到速率20

### bullet_behave_sine
水平简谐运动：ax = -omega^2*(x-0.5)，累加到vx；垂直方向不变

### bullet_behave_track_player
计算子弹到玩家的方向向量→归一化到速率5→atan2计算角度→设置LVGL图像旋转使子弹尖指向玩家→创建2秒一次性定时器，到期移除追踪行为
""",

"obj/enemy/enemy.c": """## 概述
敌人系统实现。enemy_t继承game_obj_t，基于对象池(MAX_ENEMY_COUNT=10)。包含HP管理、状态效果(灼烧/冻结)、事件响应。

### enemy_init
清零数组→初始化池→预创建LVGL图像+HP条+蓝色冻结遮罩→注册3个事件

### enemy_spawn
从池分配→设置位置/速度/血量/外观/行为→显示HP条

### enemy_move
冻结时不能移动(early return)→坐标更新→出界销毁

### 状态效果系统
- **灼烧(enemy_apply_burn)**: 不叠加，每1秒扣100HP，共2次(2000ms)
- **冻结(enemy_apply_freeze)**: 不叠加，2秒内无法移动/攻击，显示蓝色遮罩(50%透明)
- **enemy_event_hit_by_bullet_cb**: 扣血+检查子弹标志位→若BURN则施加灼烧→若FREEZE则施加冻结

### enemy_modify_hp
HP+=delta→HP<=0时派发EVENT_ENEMY_DESTROYED→调用hide归还对象池
""",

"obj/enemy/enemy_behaviors.c": """## 概述
敌人AI实现：普通敌人随机移动+周期射击，Boss两阶段循环(弹幕+追踪弹)，死亡掉落金币。

### enemy_behave_normal
v==BEHAVE_ON_DEATH时在位置生成1枚金币(值50)。首次调用时创建500ms随机移动定时器+1500ms射击定时器(发射圆形弹)

### enemy_behave_boss
v==BEHAVE_ON_DEATH时生成8枚金币(值60)。首次调用时定位到屏幕中央上方→创建300ms主时钟→两阶段：
- Phase 0 (8 ticks=2.4s): 每2 tick发射270度对称弹幕(8颗圆形弹，跳过头顶90度)
- Phase 1 (8 ticks): 每3 tick发射1颗追踪弹(三角形弹，追踪玩家2秒)
使用泰勒展开sin/cos计算弹幕方向(MCU无math.h)

### 关键宏
BEHAVE_ON_DEATH=((void*)0xDEAD) 死亡哨兵
""",

"obj/coin/coin.c": """## 概述
金币系统。coin_t继承game_obj_t，池容量13。支持玩家拾取和自动消失闪烁动画。

### coin_spawn
从池分配→设置位置/值/外观→disappear_time_s>0时创建定时器自动消失

### coin_event_hit_player_cb
玩家拾取→coin_add_num(value)→隐藏金币

### coin_disappear_timer_cb
创建LVGL动画链：透明度0→255→0，125ms半周期，重复4次(共1秒闪烁)→动画结束回调隐藏

### coin_hide
隐藏图像+标记不活跃+归还对象池+lv_anim_del停止所有动画
""",

"obj/flame_wall/flame_wall.c": """## 概述
火墙对象。Ember的Y技能，池容量3，向上移动(vy=-8)，伤害固定20。

### flame_wall_create
从池分配→设置位置→vy=-8→显示；碰撞检测在game.c中处理：火墙碰敌人扣血+隐藏火墙，火墙碰敌方子弹清除子弹

### flame_wall_move
y+=vy→出界(y<-20或y>620)时隐藏
""",
    }

    for fname, content in files.items():
        add_section(d, fname, content)

    d.save(f"{OUT}/Beat_Plane_游戏对象.docx")
    print("游戏对象.docx 已生成")


# ========================================================
#  文档4: UI界面
# ========================================================
def make_ui_docs():
    d = new_doc()
    h0(d, "Beat Plane — UI界面代码详解")

    files = {
"ui/ui.c": """## 概述
UI总调度模块。ui_init()一次初始化所有子界面，ui_run()根据FSM状态切换屏幕和音频，ui_esc_pressed_handler()为B键(ESC)的统一处理。

### ui_init
apr_init→各子界面init(9个)→注册ESC键短按回调

### ui_run (主循环持续调用)
FSM状态变化时：离开CG释放资源→switch根据新状态加载屏幕+BGM管理→更新last_game_state

### ui_esc_pressed_handler
GS_CG:跳过动画→GS_MENU / GS_PLAY→GS_PAUSE / GS_PAUSE→GS_PLAY / GS_OVER→GS_MENU / GS_SHOP:商店退出逻辑 / GS_SETTING:回到上一状态
""",

"ui/ui_play/ui_play.c": """## 概述
游戏进行中UI，项目中最复杂的UI文件。包含HUD、暂停弹窗、结束弹窗、受伤闪红、冰冻减速遮罩、技能CD可视化、Level进场动画。

### ui_play_init
创建dp_play父容器→HUD图片+金币栏→暂停/结束弹窗(含按钮)→暂停图标按钮→得分标签→受伤红色闪烁遮罩(1024x600,hurted.bin,40%透明)→冰冻减速遮罩(冰蓝半透明)→技能CD环(右侧2个lv_arc+"E"/"F"标签+"Ready!"文字)→100ms更新定时器

### ui_play_run
加载dp_play屏幕→GS_OVER:隐藏暂停按钮/显示结束弹窗/计算得分→GS_PAUSE:显示暂停弹窗→GS_PLAY:隐藏所有弹窗

### ui_play_set_freeze_overlay
控制冰冻减速遮罩的显示/隐藏(Stream Y技能)

### cd_update_timer_cb
每100ms：获取X/Y技能CD总时长和已过时间→计算百分比→更新arc值→满则显示"Ready!"否则隐藏

### hurt系统
EVENT_PLAYER_HIT_ENEMY/EVENT_BULLET_HIT_PLAYER→显示红色遮罩→创建200ms一次性定时器→到期隐藏
""",

"ui/ui_menu/ui_menu.c": """## 概述
主菜单界面。4个入口按钮(Level/Shop/Base/Setting)，右上角齿轮快速进入设置。

### ui_menu_init
创建dp_menu→背景图→Level/Shop/Base三个纵向按钮(200x90)→右上角透明Setting按钮

### ui_menu_run
加载屏幕+设置输入组焦点(用于摇杆导航)

### 按钮回调
Level→GS_PLAY+EVENT_GAME_START / Shop→GS_SHOP / Base→GS_BASE / Settings→记录上一状态后切GS_SETTING
""",

"ui/ui_shop/ui_shop.c": """## 概述
商店抽奖界面。3x3九宫格抽奖盘，加权随机概率算法+跑马灯阻尼动画+保底机制。

### ui_shop_init
创建dp_shop→背景图→金币显示→8个外围奖池格子(含图片+文字标签)→中心Draw按钮(显示消耗160金币)→跑马灯高亮光标→中奖弹窗+退出确认弹窗

### draw_btn_event_cb
校验(动画中/弹窗未关/金币不足则拒)→扣160金币→get_random_reward_slot确定目标→计算步数(最少3圈+到目标)→创建定时器启动跑马灯

### roulette_timer_cb
每次步进1格(共8格循环)→距离目标<15步时增加间隔实现阻尼减速→步数耗尽后give_reward+show_reward_popup

### get_random_reward_slot
第1次固定出Stream(格2)，第6次固定出Verdant(格4)；其余基于权重万分比随机

### give_reward
根据格子增加金币或调用ui_base_plane_unlock解锁飞机
""",

"ui/ui_base/ui_base.c": """## 概述
基地/机库选飞机界面。4架飞机横向展示，锁定/解锁状态，详情面板，飞机选择逻辑。

### ui_base_init
创建dp_base→背景图→4个飞机卡片(图片+名字+锁遮罩，横向排列)→详情面板(HP/DMG/Skill 3行+Choose按钮)→退出确认弹窗

### ui_base_run
加载屏幕→根据解锁状态显示/隐藏锁遮罩→默认展示当前选中飞机的详情面板→对齐SELECTED标签

### 飞机模板
- Player: "E:3-Way Shot / F:Shield 1s"
- Ember: "E:Burn Bullet / F:Flame Wall"
- Stream: "E:Freeze Bullet / F:Slow Enemy"
- Verdant: "E:Speed Boost / F:HP Reclaim"

### choose_btn_cb
校验解锁状态→未解锁则拒绝→解锁则更新g_selected_plane_id
""",

"ui/ui_setting/ui_setting.c": """## 概述
设置界面。音量调节(BGM/SFX/放大器3个滑块+预览按钮)、Hitbox碰撞框可视化开关、清除存档功能。

### ui_setting_init
创建dp_setting→滚动容器→Amp滑块(0-3档:50%/100%/200%/400%)→BGM滑块(0-100→内部0-255)+预览按钮→SFX滑块(同)→Hitbox开关+飞机预览图+红色碰撞框→CLEAR SAVE DATA按钮+确认弹窗→返回按钮

### 滑块回调
将UI值(0-100)映射到内部0-255调用audio_set_vol_xxx()；释放滑块时播放预览音频

### clear_save_btn_event_cb
显示确认弹窗(隐藏主容器)→确认→save_clear()+save_load()→取消→恢复主容器
""",

"ui/ui_cg/ui_cg.c": """## 概述
开场CG动画。LVGL动画链实现图片/文字的淡入淡出序列，可点击跳过。

### ui_cg_run (动画时间轴)
0-5s: label1显示 → 5s: label1淡出(300ms) → 5s: img1淡入(800ms) → 10s: img1淡出(400ms) → 11s: img2+label2淡入(800ms) → 16s: 全部淡出(500ms) → 16.5s: cg_layer淡出(1500ms) → 完成→500ms后切GS_MENU

### ui_cg_skip
删除所有动画和定时器→停止BGM→切GS_MENU

### ui_cg_cleanup
释放cg_layer和图片描述符内存(模拟器环境)
""",

"ui/ui_key/ui_key.c": """## 概述
硬件按键到LVGL输入设备的桥接。摇杆→LV_KEY_PREV/NEXT，A键→LV_KEY_ENTER。

### key_get_indev
单例模式创建LVGL键盘输入设备

### ui_key_read_cb (LVGL自动调用)
GS_PLAY时返回REL(游戏中进行中禁用UI导航)→摇杆上/左→LV_KEY_PREV→摇杆下/右→LV_KEY_NEXT→A键→LV_KEY_ENTER
""",

"ui/ui_sys_halt/ui_sys_halt.c": """## 概述
系统停机界面。不可恢复错误时显示"SYSTEM HALT"+日志列表。真机关中断死循环，模拟器5秒后exit(-1)。

### ui_sys_halt_run
加载屏幕→创建日志标签(Error红色/Warning橙色/其他绿色)→lv_timer_handler立即刷新→真机:__disable_irq()+while(1)→模拟器:5秒后exit(-1)
""",

"ui/ui_templates/ui_templates.c": """## 概述
UI可复用模板。仅实现popup_create()函数。

### popup_create
创建400x500圆角(20px)弹窗→居中→背景色0x121212(70%不透明)→边框色0x4FC3F7(宽度3)→移到前景→去除滚动
""",
    }

    for fname, content in files.items():
        add_section(d, fname, content)

    d.save(f"{OUT}/Beat_Plane_UI界面.docx")
    print("UI界面.docx 已生成")


# ========================================================
#  文档5: 驱动与工具
# ========================================================
def make_driver_docs():
    d = new_doc()
    h0(d, "Beat Plane — 驱动与工具代码详解")

    files = {
"driver/audio/audio.c": """## 概述
音频系统。PC端用SDL2音频回调，MCU端用I2S+DMA。支持BGM+SFX(3个音效通道)多通道混音，带独立音量控制。

### audio_init
SDL: 打开44.1kHz 16bit单声道音频设备→SDL_PauseAudio(0)开始播放 / MCU: 配置I2S+GPIO+SPI中断

### audio_load(id, channel, repeat)
在指定通道播放音频资源。7个音频资源预定义路径和大小(cg.pcm 1.8M / bgm.pcm 13M等)

### audio_play_on_channel (static)
锁外: ram_malloc分配内存+read_file_to_array读取文件→锁内: 原子交换通道数据指针(旧数据延迟释放)→激活播放

### 混音算法 (mix_wavg)
加权平均: sum = vol_bgm*AUDIO_ALLOC_BGM*bgm + Σ(vol_sfx*AUDIO_ALLOC_SFX*sfx[i]) → 移位归一化 → 饱和截断(-32768~32767)

### 音量控制
vol_bgm/vol_sfx: 0-255；vol_amp: 0-3档(对应0.5x/1x/2x/4x)
""",

"driver/input_hw/key/key.c": """## 概述
按键虚拟化系统。将物理按键(SDL键码/GPIO引脚)抽象为KEY_A/B/X/Y虚拟按键。

### key_init
PC: 无操作 / MCU: 配置GPIO上拉输入模式

### key_scan (每5ms)
遍历所有按键→key_process(读取+消抖+边沿检测+长按检测)

### key_process (核心)
1.读取原始状态→2.MCU端5ms消抖→3.边沿检测(pressed=上升沿/released=下降沿)→4.长按检测(按下持续>LONG_PRESS_MS=150ms)→5.同步last=stable

### PC按键绑定
- KEY_A: 空格+回车 (射击)
- KEY_B: Q+ESC (返回)
- KEY_X: E (X技能)
- KEY_Y: F (Y技能)

### MCU按键绑定
- KEY_A: S3按钮 / KEY_B: S4 / KEY_X: S5 / KEY_Y: S2
""",

"driver/input_hw/joystick/joystick.c": """## 概述
摇杆输入。PC端WASD+方向键模拟，MCU端ADC采集。

### joystick_init
PC: 无 / MCU: 配置ADC

### joystick_scan (每5ms)
PC: 读取SDL键盘状态→加速衰减滤波→输出-127~127 / MCU: 读取ADC→中值校准→死区处理

### joystick_get_x / joystick_get_y
返回当前摇杆值(-256~256范围，config.h可配)
""",

"driver/lv_port/lv_port.c": """## 概述
LVGL显示移植。PC端创建SDL窗口(1024x600, LV_COLOR_DEPTH=16)，MCU端配置LCD控制器。

### lv_port_init
PC: 创建SDL窗口+双缓冲(1/10屏幕高)→注册flush回调(SDL_UpdateTexture+SDL_RenderCopy+SDL_RenderPresent) / MCU: 配置LTDC+层混合+DMA2D
""",

"driver/uart/uart.c": """## 概述
UART串口通信。PC端使用Windows COM端口API(CreateFile/ReadFile/WriteFile)，MCU端使用USART硬件。

### uart_enable
PC: 打开COM口 / MCU: 使能USART时钟+GPIO复用+中断

### uart_send
发送字节数组

### uart_recv
接收字节(带超时)
""",

"tools/tools.c": """## 概述
基础工具函数集。

### play_tick_get
播放时钟：仅在GS_PLAY时累加(暂停/菜单时冻结)，返回累计毫秒数。用于技能CD和时间相关逻辑。

### non_blocking_delay
非阻塞定时器检查：若距离上次调用>=delay_ms则调用func

### ram_malloc / ram_free
内存管理封装(MCU端可能自定义堆)

### direction_to_velocity
根据方向向量(dx,dy)和目标速率计算vx,vy

### vec_length
向量长度(用于归一化)
""",

"tools/lvgl_utils/lvgl_utils.c": """## 概述
LVGL辅助工具。图片路径构建、文件读取、图片描述符加载/释放。

### img_path
构建图片完整路径：SIMULATOR="./assets/pics/" + name / MCU="0:/assets/pics/" + name

### read_file_to_array
文件→内存缓冲。SIMULATOR用fopen/fread，MCU用FatFS f_open/f_read。
EMBED_ASSETS模式下先查嵌入数据，命中则memcpy(不读磁盘)→未命中回退到文件读取。

### load_img_dsc
读取图片文件→分配内存→填充lv_img_dsc_t.header(宽/高/颜色格式)→设置data指针→alpha图片偏移4字节跳过头部

### free_img_dsc
释放图片描述符内存(alpha图片回退4字节)

### popup_show / popup_hide
清除/添加LV_OBJ_FLAG_HIDDEN

### set_group
设置输入设备目标组
""",

"tools/ring_buffer/ring_buffer.c": """## 概述
环形缓冲区(线程安全)。双指针+锁实现FIFO。

### ring_buf_init
设置buffer/size→清零读写指针

### ring_buf_write
写数据到缓冲区，写指针前进

### ring_buf_read
从缓冲区读数据，读指针前进

### ring_buf_available
返回可读字节数
""",

"tools/perf_monitor/perf_monitor.c": """## 概述
性能监视器。显示每帧逻辑耗时(mspt)的UI标签。

### perf_monitor_init
创建半透明标签(左下角)，显示帧时间

### perf_monitor_update
更新标签文本

### perf_monitor_set_mspt
由game_update调用记录每帧耗时
""",

"tools/debug/debug.c": """## 概述
调试系统。提供调试菜单和可视化工具(碰撞箱显示等)。
""",

"comm/comm.c": """## 概述
通信层入口。统一通信接口，PC/MCU双端。

### comm_init
初始化通信子系统(PC:COM端口 / MCU:UART)

### comm_update
通信主循环(500Hz)，处理收发
""",

"comm/protocol.h": """## 概述
通信协议定义。帧格式：帧头(0xAA 0x55)+指令码+数据长度+数据+校验。

## 关键宏
COMM_KEY_A_MASK(0x01) / COMM_KEY_B_MASK(0x02) / COMM_KEY_X_MASK(0x04) / COMM_KEY_Y_MASK(0x08)
""",
    }

    for fname, content in files.items():
        add_section(d, fname, content)

    d.save(f"{OUT}/Beat_Plane_驱动与工具.docx")
    print("驱动与工具.docx 已生成")


# ========================================================
if __name__ == "__main__":
    make_obj_docs()
    make_ui_docs()
    make_driver_docs()
    print("All done!")
