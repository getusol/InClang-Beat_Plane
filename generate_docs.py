"""
generate_docs.py — 从 agent 分析结果生成 Word 文档
用法: python generate_docs.py
"""
import json, os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT_DIR = "docs"
TASK_DIR = r"C:\Users\DENGXI~1\AppData\Local\Temp\claude\t--github-new-Beat-Plane\af43dcc0-0b16-44f4-940f-0ae7098f3824\tasks"

# 4 个 agent 的输出文件
TASK_FILES = {
    "core":   "a36aa9d667b68e899.output",
    "obj":    "a41267af11a35858a.output",
    "ui":     "a2d39a6834e63dfeb.output",
    "driver": "ac109584435e7eb08.output",
}

def read_agent_output(filename):
    """读取 JSONL 格式的 agent 输出，提取最终文本"""
    path = os.path.join(TASK_DIR, filename)
    if not os.path.exists(path):
        print(f"WARNING: {path} not found")
        return ""
    text_parts = []
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            try:
                obj = json.loads(line)
                if obj.get("type") == "assistant" and "text" in obj:
                    text_parts.append(obj["text"])
            except:
                pass
    return "\n".join(text_parts)

def parse_sections(text):
    """将 agent 输出按 '=== 文件:' 分割成段落"""
    sections = re.split(r'(?=== 文件: )', text)
    result = []
    for s in sections:
        s = s.strip()
        if not s:
            continue
        # 提取文件名
        m = re.match(r'=== 文件:\s*(.+?)\s*===', s)
        if m:
            fname = m.group(1).strip()
            body = s[m.end():].strip()
            # 清理路径前缀
            fname = fname.replace("t:\\github_new\\Beat_Plane\\src\\", "")
            fname = fname.replace("t:/github_new/Beat_Plane/src/", "")
            fname = fname.replace("\\", "/")
            result.append((fname, body))
    return result

def add_code_block(doc, code_text):
    """添加带格式的代码块"""
    for line in code_text.split("\n"):
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def create_per_file_docs(all_sections):
    """为每个分类创建 Word 文档"""
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # 分类
    categories = {
        "01_主入口与配置": lambda f: "main.c" in f or "config.h" in f,
        "02_游戏核心": lambda f: any(x in f for x in ["game/game.", "core/fsm", "core/event", "core/apr", "core/game_object", "core/pool", "core/timer", "core/input_sw", "level/", "save/"]),
        "03_游戏对象": lambda f: any(x in f for x in ["obj/player", "obj/bullet", "obj/enemy", "obj/coin", "obj/flame"]),
        "04_UI界面": lambda f: any(x in f for x in ["ui/ui_", "ui/ui."]),
        "05_驱动层": lambda f: any(x in f for x in ["driver/", "comm/"]),
        "06_工具层": lambda f: any(x in f for x in ["tools/"]),
    }

    for cat_name, cat_filter in categories.items():
        doc = Document()
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = "微软雅黑"
        style.font.size = Pt(10)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 标题
        title = doc.add_heading(f"Beat Plane — {cat_name}", level=0)

        cat_sections = [(f, b) for f, b in all_sections if cat_filter(f)]
        if not cat_sections:
            print(f"  No sections for {cat_name}")
            continue

        for fname, body in sorted(cat_sections):
            doc.add_heading(fname, level=1)

            # 解析 body 的子段
            parts = re.split(r'\n(?=## )', body)
            for part in parts:
                part = part.strip()
                if not part:
                    continue

                # 二级标题
                h2m = re.match(r'##\s+(.+)', part)
                if h2m:
                    h2text = h2m.group(1)
                    doc.add_heading(h2text, level=2)
                    rest = part[h2m.end():].strip()
                else:
                    rest = part

                # 处理三级标题和代码块
                lines = rest.split("\n")
                i = 0
                in_code = False
                code_buf = []

                while i < len(lines):
                    line = lines[i]

                    # 三级标题
                    if line.startswith("### "):
                        if in_code:
                            add_code_block(doc, "\n".join(code_buf))
                            code_buf = []
                            in_code = False
                        h3text = line[4:].strip()
                        doc.add_heading(h3text, level=3)
                        i += 1
                        continue

                    # 关键代码标记
                    if "**关键代码**:" in line:
                        if in_code:
                            add_code_block(doc, "\n".join(code_buf))
                            code_buf = []
                            in_code = False
                        i += 1
                        continue

                    # 代码块开始/结束
                    if line.strip() == "```c":
                        if in_code:
                            add_code_block(doc, "\n".join(code_buf))
                            code_buf = []
                            in_code = False
                        else:
                            in_code = True
                        i += 1
                        continue

                    if in_code:
                        code_buf.append(line)
                    else:
                        # 普通文本段落
                        stripped = line.strip()
                        if stripped and not stripped.startswith("#"):
                            # 处理 **粗体** 标记
                            p = doc.add_paragraph()
                            p.paragraph_format.space_after = Pt(3)
                            # 简单处理 bold
                            parts_bold = re.split(r'(\*\*.+?\*\*)', stripped)
                            for pb in parts_bold:
                                if pb.startswith("**") and pb.endswith("**"):
                                    run = p.add_run(pb[2:-2])
                                    run.bold = True
                                else:
                                    p.add_run(pb)

                    i += 1

                if in_code and code_buf:
                    add_code_block(doc, "\n".join(code_buf))

        path = os.path.join(OUTPUT_DIR, f"Beat_Plane_{cat_name}.docx")
        doc.save(path)
        print(f"  Created: {path} ({len(cat_sections)} files)")

def create_overview_doc(all_sections):
    """创建总览文档"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = "微软雅黑"
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_heading("Beat Plane 项目代码总览", level=0)

    # 项目架构概述
    doc.add_heading("一、项目架构", level=1)
    doc.add_paragraph(
        "Beat Plane 是一款基于 LVGL (LittlevGL) 图形库的 2D 飞机射击游戏，"
        "支持 PC 模拟器 (SDL2) 和 STM32 单片机双平台运行。\n\n"
        "核心架构分为 6 层："
    )

    layers = [
        ("主入口层 (main.c)", "主循环使用 4 个非阻塞定时器以不同频率驱动：输入扫描(5ms)、游戏逻辑(30Hz)、UI刷新(30Hz)、通信(500Hz)。"),
        ("游戏核心层 (game/)", "有限状态机(FSM)管理游戏全局状态；事件系统提供发布-订阅解耦；对象池管理有限资源；外观系统(APR)集中管理视觉模板；定时器系统支持一次性/重复回调。"),
        ("游戏对象层 (obj/)", "所有实体继承 game_obj_t 基类，包含玩家(4种飞机可切换)、子弹(池容量50)、敌人(池容量10)、金币(池容量13)、火墙(池容量3)。每种对象有独立的 update/show/hide 生命周期。"),
        ("UI界面层 (ui/)", "9个界面模块：CG开场动画、主菜单、游戏进行(HUD+暂停+结束)、设置(音量/碰撞框)、商店(抽奖)、基地(选机)、通信状态、按键输入桥接、系统停机。"),
        ("驱动层 (driver/)", "硬件抽象：音频(I2S/SDL)、输入(键盘/摇杆/GPIO)、显示(LVGL移植)、UART通信。通过 #ifdef SIMULATOR 宏实现PC/MCU双平台兼容。"),
        ("工具层 (tools/)", "基础工具函数、LVGL辅助(图片加载/弹窗)、环形缓冲区、调试系统、性能监视器、存档系统。"),
    ]

    for name, desc in layers:
        doc.add_heading(name, level=2)
        doc.add_paragraph(desc)

    # 数据流
    doc.add_heading("二、游戏主循环数据流", level=1)
    doc.add_paragraph(
        "main() → while(1):\n"
        "  1. input_dispatch() — 扫描按键/摇杆，分发给注册的回调 (短按/长按/持续按下)\n"
        "  2. game_update() — 更新定时器 → 遍历活跃对象调用 update/behave → AABB碰撞检测 → 关卡更新\n"
        "  3. ui_run() — FSM状态变化时切换屏幕和音频\n"
        "  4. comm_update() — PC/MCU 间双向通信\n"
        "  5. lv_timer_handler() — LVGL 内部渲染和定时器处理\n"
        "  6. delay_ms(1)\n\n"
        "游戏逻辑帧率: GAME_TICK=30Hz，即每约33ms执行一次 game_update()"
    )

    # 状态机流转
    doc.add_heading("三、FSM 状态流转", level=1)
    states = [
        "GS_CG (开场动画) → 自动播放完毕后 → GS_MENU",
        "GS_MENU → Level按钮 → GS_PLAY",
        "GS_MENU → Shop按钮 → GS_SHOP",
        "GS_MENU → Base按钮 → GS_BASE",
        "GS_PLAY → 暂停键 → GS_PAUSE → 继续键 → GS_PLAY",
        "GS_PLAY → 玩家死亡 → GS_OVER → Restart → GS_PLAY",
        "GS_OVER → Back to menu → GS_MENU",
        "GS_SHOP → 返回 → GS_MENU",
        "GS_BASE → 返回 → GS_MENU",
        "SYS_HALT — 不可恢复的停机状态",
    ]
    for s in states:
        doc.add_paragraph(s, style='List Bullet')

    # 文件清单
    doc.add_heading("四、文件清单 (56个源文件)", level=1)

    file_groups = {
        "主入口": ["main.c", "config.h"],
        "游戏核心": ["game/game.c", "game/game.h", "core/fsm/fsm.c", "core/fsm/fsm.h",
                   "core/event/event.c", "core/event/event.h", "core/apr/apr.c", "core/apr/apr.h",
                   "core/game_object/game_object.c", "core/game_object/game_object.h",
                   "core/pool/pool.c", "core/pool/pool.h", "core/timer/timer.c", "core/timer/timer.h",
                   "core/input_sw/input_sw.c", "core/input_sw/input_sw.h",
                   "level/level.c", "level/level.h", "save/save.c", "save/save.h"],
        "游戏对象": ["obj/player/player.c", "obj/player/player.h",
                   "obj/bullet/bullet.c", "obj/bullet/bullet.h",
                   "obj/bullet/bullet_behaviors.c", "obj/bullet/bullet_behaviors.h",
                   "obj/enemy/enemy.c", "obj/enemy/enemy.h",
                   "obj/enemy/enemy_behaviors.c", "obj/enemy/enemy_behaviors.h",
                   "obj/coin/coin.c", "obj/coin/coin.h",
                   "obj/flame_wall/flame_wall.c", "obj/flame_wall/flame_wall.h"],
        "UI界面": ["ui/ui.c", "ui/ui.h", "ui/ui_play/ui_play.c", "ui/ui_play/ui_play.h",
                  "ui/ui_base/ui_base.c", "ui/ui_base/ui_base.h",
                  "ui/ui_menu/ui_menu.c", "ui/ui_menu/ui_menu.h",
                  "ui/ui_shop/ui_shop.c", "ui/ui_shop/ui_shop.h",
                  "ui/ui_setting/ui_setting.c", "ui/ui_setting/ui_setting.h",
                  "ui/ui_cg/ui_cg.c", "ui/ui_cg/ui_cg.h",
                  "ui/ui_comm/ui_comm.c", "ui/ui_comm/ui_comm.h",
                  "ui/ui_key/ui_key.c", "ui/ui_key/ui_key.h",
                  "ui/ui_sys_halt/ui_sys_halt.c", "ui/ui_sys_halt/ui_sys_halt.h",
                  "ui/ui_templates/ui_templates.c", "ui/ui_templates/ui_templates.h"],
        "驱动层": ["driver/audio/audio.c", "driver/audio/audio.h",
                  "driver/input_hw/input_hw.c", "driver/input_hw/input_hw.h",
                  "driver/input_hw/key/key.c", "driver/input_hw/key/key.h",
                  "driver/input_hw/joystick/joystick.c", "driver/input_hw/joystick/joystick.h",
                  "driver/lv_port/lv_port.c", "driver/lv_port/lv_port.h",
                  "driver/uart/uart.c", "driver/uart/uart.h"],
        "工具层": ["tools/tools.c", "tools/tools.h",
                  "tools/lvgl_utils/lvgl_utils.c", "tools/lvgl_utils/lvgl_utils.h",
                  "tools/ring_buffer/ring_buffer.c", "tools/ring_buffer/ring_buffer.h",
                  "tools/debug/debug.c", "tools/debug/debug.h",
                  "tools/pref_monitor/perf_monitor.c", "tools/pref_monitor/perf_monitor.h"],
        "通信层": ["comm/comm.c", "comm/comm.h", "comm/comm_rx.c", "comm/comm_rx.h",
                  "comm/comm_tx.c", "comm/comm_tx.h", "comm/comm_status.c", "comm/comm_status.h",
                  "comm/protocol.h"],
    }

    for group, files in file_groups.items():
        doc.add_heading(group, level=2)
        for f in files:
            doc.add_paragraph(f"• {f}")

    path = os.path.join(OUTPUT_DIR, "Beat_Plane_总览.docx")
    doc.save(path)
    print(f"  Created: {path}")

# ====== 主流程 ======
if __name__ == "__main__":
    print("Reading agent outputs...")
    all_sections = []
    for key, filename in TASK_FILES.items():
        text = read_agent_output(filename)
        if text:
            sections = parse_sections(text)
            print(f"  {key}: {len(sections)} file sections parsed")
            all_sections.extend(sections)

    print(f"\nTotal: {len(all_sections)} file sections")

    print("\nGenerating overview doc...")
    create_overview_doc(all_sections)

    print("\nGenerating per-category docs...")
    create_per_file_docs(all_sections)

    print("\nDone! Check the docs/ folder.")
