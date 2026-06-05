"""
生成 Beat Plane 代码文档 Word 文件
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn

OUT = "docs"
os.makedirs(OUT, exist_ok=True)

def new_doc():
    d = Document()
    s = d.styles['Normal']
    s.font.name = "Consolas"
    s.font.size = Pt(9)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return d

def h0(d, t): d.add_heading(t, level=0)
def h1(d, t): d.add_heading(t, level=1)
def h2(d, t): d.add_heading(t, level=2)
def h3(d, t): d.add_heading(t, level=3)
def para(d, t):
    p = d.add_paragraph(t)
    p.paragraph_format.space_after = Pt(4)
def bullet(d, t): d.add_paragraph(t, style='List Bullet')
def code(d, t):
    for line in t.split("\n"):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(line)
        r.font.size = Pt(7); r.font.name = "Consolas"

# ============================================================
#  总览文档
# ============================================================
def make_overview():
    d = new_doc()
    h0(d, "Beat Plane 项目代码总览")

    h1(d, "一、项目简介")
    para(d, "Beat Plane 是一款基于 LVGL (LittlevGL v8.2) 图形库的 2D 飞机射击游戏。"
         "支持 PC 模拟器 (SDL2) 和 STM32 单片机双平台运行。"
         "代码通过 #ifdef SIMULATOR 宏实现同一套游戏逻辑在两个平台上的无缝切换。")

    h1(d, "二、架构分层 (6层)")
    layers = [
        ("主入口 (main.c)", "while(1) 主循环，4个非阻塞定时器以不同频率驱动各子系统：input(5ms)→game_update(30Hz)→ui_run(30Hz)→comm(500Hz)→lv_timer_handler()→delay_ms(1)",
         "config.h 定义所有全局常量：屏幕1024x600、子弹上限50、敌人上限10、游戏帧率30Hz等"),
        ("游戏核心 (game/core/)", "FSM状态机管理 CG→Menu→Play(Pause/Over)→Shop→Base 全状态流转；"
         "事件系统(发布-订阅)解耦碰撞检测与业务回调；"
         "对象池(pool)预分配固定数量槽位，用 O(1) 空闲栈管理分配/回收；"
         "APR外观系统用枚举+查表统一管理所有视觉模板(尺寸/图片/碰撞箱)；"
         "定时器(timer)支持 ONCE/REPEAT 模式，自动跟随游戏对象的生命周期"),
        ("游戏对象 (game/obj/)", "player: 4种飞机可切换(Player/Ember/Stream/Verdant)，各有X/Y两个主动技能；"
         "bullet: 对象池50发，支持圆周/正弦/追踪3种行为，减速/灼烧/冻结3种效果标志；"
         "enemy: 对象池10个，普通敌人随机移动+射击，Boss两阶段(270度弹幕+追踪弹)；"
         "coin: 对象池13枚，拾取增加金币，支持自动消失闪烁动画；"
         "flame_wall: 对象池3个，Ember火墙技能，向上移动清除路径上的敌方子弹并伤害敌人"),
        ("UI界面 (ui/)", "ui.c 为总调度，根据FSM状态切换屏幕和BGM；"
         "ui_play: 游戏HUD(HUD+HP条+金币+暂停/结束弹窗+受伤闪红+冰冻遮罩+技能CD环)；"
         "ui_menu: 主菜单(Level/Shop/Base/Setting 4按钮)；"
         "ui_shop: 九宫格抽奖(加权随机+跑马灯动画+保底机制)；"
         "ui_base: 机库选飞机(4架横向展示+详情面板+锁定/解锁)；"
         "ui_setting: 音量调节/BGM预览/Hitbox开关/清档；"
         "ui_cg: 开场CG动画链(淡入淡出序列+可跳过)；"
         "ui_key: 摇杆→LVGL按键导航桥接；ui_comm: 通信状态指示"),
        ("驱动层 (driver/)", "audio: I2S/SDL音频，支持BGM+SFX多通道混音，带音量控制；"
         "input_hw: 按键扫描(去抖/长按检测)，摇杆读取(ADC/SDL)；"
         "key: 虚拟按键系统，PC端绑定SDL键码(空格/回车→A, Q/ESC→B, E→X, F→Y)，MCU端绑定GPIO(S2-S5)；"
         "joystick: PC端WASD+方向键模拟摇杆，MCU端ADC采集；"
         "lv_port: LVGL显示移植(SDL窗口/MCU液晶屏)；"
         "uart: PC端COM端口，MCU端USART硬件"),
        ("工具层 (tools/)+通信层 (comm/)", "tools: play_tick_get(仅GS_PLAY递增的计时器)、非阻塞定时器、方向归一化；"
         "lvgl_utils: 图片路径构建、文件读取(read_file_to_array)、图片描述符加载/释放、弹窗管理；"
         "ring_buffer: 线程安全的环形缓冲区(双指针+锁)；"
         "debug: 调试菜单+可视化；perf_monitor: 帧性能统计UI；save: key=value格式存档(支持自动备份)；"
         "comm: PC↔MCU双向通信协议(帧头+指令+长度+数据+校验)、连接状态机")
    ]
    for name, desc, *extra in layers:
        h2(d, name)
        para(d, desc)
        if extra: para(d, extra[0])

    h1(d, "三、游戏主循环数据流")
    para(d, "main() 初始化所有子系统后进入 while(1)，每轮执行：\n"
         "1. non_blocking_delay(&input_timer) → input_dispatch() 扫描按键/摇杆\n"
         "2. non_blocking_delay(&logic_timer) → game_update() 游戏逻辑(30Hz)\n"
         "   ├─ timer_update() 更新所有游戏定时器\n"
         "   ├─ for each active obj: obj→update(obj) → obj→behave.f(obj) → hitbox_update\n"
         "   ├─ check_collisions() AABB碰撞检测，按类型组合派发事件\n"
         "   └─ level_update() 关卡波次管理\n"
         "3. non_blocking_delay(&ui_timer) → ui_run() FSM状态切换\n"
         "4. non_blocking_delay(&comm_timer) → comm_update() 通信处理\n"
         "5. lv_timer_handler() LVGL内部渲染和定时器\n"
         "6. perf_monitor_update(); delay_ms(1)")

    h1(d, "四、FSM 状态流转图")
    for s in ["GS_CG → (动画完/跳过) → GS_MENU",
              "GS_MENU → [Level] → GS_PLAY | [Shop] → GS_SHOP | [Base] → GS_BASE | [Settings] → GS_SETTING",
              "GS_PLAY → [暂停] → GS_PAUSE → [继续] → GS_PLAY | [玩家死亡] → GS_OVER",
              "GS_OVER → [Restart] → GS_PLAY | [Back] → GS_MENU",
              "GS_SHOP/GS_BASE → [返回] → GS_MENU",
              "GS_SETTING → [返回] → 回到上一状态 (MENU/PLAY/PAUSE)",
              "SYS_HALT ← 不可恢复的错误(无效状态/内存失败)"]:
        bullet(d, s)

    h1(d, "五、对象池设计模式")
    para(d, "所有游戏实体(子弹/敌人/金币/火墙/定时器)都使用对象池管理：\n"
         "• 启动时预分配固定数量槽位(静态数组)\n"
         "• 空闲栈(free_stack)用 uint16_t 数组维护可用索引\n"
         "• pool_alloc(): O(1) 弹出栈顶索引\n"
         "• pool_free(): O(1) 压回索引(防重复归还)\n"
         "• 最大容量由 config.h 宏控制，MCU上无动态内存分配")

    h1(d, "六、事件系统")
    para(d, "11种事件类型：GAME_START, PLAYER_HIT_ENEMY, BULLET_HIT_ENEMY, BULLET_HIT_PLAYER, "
         "PLAYER_DIE, ENEMY_DESTROYED, PLAYER_HIT_COIN, FLAME_WALL_HIT_ENEMY 等。\n"
         "每个事件最多注册 MAX_EVENT_LISTENER(5) 个回调。\n"
         "碰撞检测(check_collisions)根据对象类型组合自动派发事件，各模块注册自己的回调处理业务逻辑。")

    h1(d, "七、4架飞机技能表")
    planes = [
        ("Player (id=0)", "HP200, 射速200ms, 伤害34", "X/E: Triple Shot 三向散射(3s)", "Y/F: Shield 护盾无敌1秒(5s)"),
        ("Ember (id=1)", "HP200, 射速150ms, 伤害80", "X/E: Burn Bullet 灼烧弹DOT(3s)", "Y/F: Flame Wall 火墙清除弹幕(5s)"),
        ("Stream (id=2)", "HP200, 射速250ms, 伤害10", "X/E: Freeze Bullet 冻结弹(3s)", "Y/F: Bullet Slow 全屏减速+冰霜渲染(6s)"),
        ("Verdant (id=3)", "HP250, 射速220ms, 伤害15", "X/E: Speed Boost 长按E加速(无CD)", "Y/F: HP Reclaim 回血+50(6s)"),
    ]
    for name, stats, xskill, yskill in planes:
        h2(d, name)
        para(d, f"{stats}")
        bullet(d, xskill)
        bullet(d, yskill)

    h1(d, "八、文件清单 (56个源文件)")
    groups = {
        "主入口": "main.c, config.h",
        "游戏核心(10对)": "game.c/h, fsm.c/h, event.c/h, apr.c/h, game_object.c/h, pool.c/h, timer.c/h, input_sw.c/h, level.c/h, save.c/h",
        "游戏对象(7对)": "player, bullet(+behaviors), enemy(+behaviors), coin, flame_wall (各含.c/.h)",
        "UI界面(11对)": "ui, ui_play, ui_menu, ui_shop, ui_base, ui_setting, ui_cg, ui_comm, ui_key, ui_sys_halt, ui_templates",
        "驱动(6对)": "audio, input_hw, key, joystick, lv_port, uart",
        "工具(5对)": "tools, lvgl_utils, ring_buffer, debug, perf_monitor",
        "通信(4对+protocol)": "comm, comm_rx, comm_tx, comm_status, protocol.h",
    }
    for g, files in groups.items():
        h2(d, g)
        para(d, files)

    d.save(f"{OUT}/Beat_Plane_总览.docx")
    print("总览.docx 已生成")

# ============================================================
#  分文件文档
# ============================================================
def add_section(d, fname, content):
    """向文档添加一个文件的分析"""
    d.add_heading(fname, level=1)

    # 按 ## 分割为子段
    parts = content.split("\n## ")
    for part in parts:
        part = part.strip()
        if not part: continue
        # 第一个 part 是概述 (不以 ## 开头)
        if not part.startswith("##") and "概述" in part.lower():
            h2(d, "概述")
            para(d, part.replace("## 概述", "").strip())
            continue

        lines = part.split("\n")
        i = 0
        in_code = False
        code_buf = []

        while i < len(lines):
            line = lines[i]

            # 三级标题 ###
            if line.startswith("### "):
                if in_code:
                    code(d, "\n".join(code_buf)); code_buf = []; in_code = False
                h3(d, line[4:].strip())
                i += 1; continue

            # ```c 代码块
            if line.strip() == "```c":
                if in_code:
                    code(d, "\n".join(code_buf)); code_buf = []; in_code = False
                else:
                    in_code = True
                i += 1; continue

            if in_code:
                code_buf.append(line)
            else:
                stripped = line.strip()
                if stripped and not stripped.startswith("#"):
                    # 处理 **粗体**
                    import re
                    p = d.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    parts2 = re.split(r'(\*\*.+?\*\*)', stripped)
                    for pb in parts2:
                        if pb.startswith("**") and pb.endswith("**"):
                            r = p.add_run(pb[2:-2]); r.bold = True
                        else:
                            p.add_run(pb)
            i += 1

        if in_code and code_buf:
            code(d, "\n".join(code_buf))


# ---- 游戏核心文件 (来自 Agent 1 分析) ----
def make_core_docs():
    d = new_doc()
    h0(d, "Beat Plane — 游戏核心代码详解")

    files = {
        "game/game.c": """## 概述
游戏核心管理器，负责游戏对象的注册、每帧更新遍历、碰撞检测与派发。

### game_init
- **作用**: 初始化所有游戏子系统和游戏对象
- **逻辑流程**: 获取play显示对象 → 初始化timer/level → 依次初始化player/bullet/enemy/coin/flame_wall → 注册UI事件 → 所有对象初始化碰撞箱
- **关键代码**:
```c
void game_init() {
    lv_obj_t *play_display = ui_play_get_display();
    timer_pool_init();
    level_init();
    player_init(play_display);
    bullet_init(play_display);
    enemy_init(play_display);
    coin_init(play_display);
    flame_wall_init(play_display);
    ui_play_register_events();
    game_for_each_obj(init_hitbox, NULL);
}
```

### game_update
- **作用**: 主游戏更新循环 (30Hz)，被main的logic_timer驱动
- **逻辑流程**: 更新定时器 → 遍历活跃对象调用update+behave → 更新碰撞箱 → AABB碰撞检测 → 关卡更新 → 记录性能数据
- **关键代码**:
```c
void game_update() {
    timer_update();
    for (int i = 0; i < free_idx; i++) {
        if (!game_obj_is_active(game_objs[i])) continue;
        if (game_objs[i]->update) game_objs[i]->update(game_objs[i]);
        if (game_objs[i]->behave.f) game_objs[i]->behave.f(game_objs[i], game_objs[i]->behave.usr_data);
        game_obj_hitbox_update(game_objs[i]);
    }
    check_collisions();
    level_update();
}
```

### check_collisions (static)
- **作用**: AABB碰撞检测，双层循环遍历所有对象对，按类型组合派发事件
- **逻辑流程**: 只检测特定类型组合(PLAYER-ENEMY/BULLET-ENEMY/PLAYER-COIN/PLAYER-BULLET/FLAME_WALL-ENEMY/FLAME_WALL-BULLET)；重叠时分别处理(护盾免疫/反射、火墙清弹、子弹扣血等)
""",

        "core/fsm/fsm.c": """## 概述
有限状态机，管理游戏的全局状态流转：CG→MENU→PLAY(PAUSE/OVER)→SHOP→BASE→SYS_HALT。

### fsm_init
将初始状态设为 GS_CG

### fsm_get_state
返回 current_state

### fsm_switch_state
校验状态范围 → 设置 current_state；无效状态触发 sys_halt()
""",

        "core/event/event.c": """## 概述
事件系统，发布-订阅模式。11种事件类型，每种最多5个监听者。

### event_init
清空所有回调槽位

### event_register
在指定事件的回调数组中找空位注册；防重复注册

### event_dispatch
遍历该事件的所有回调并依次调用
""",

        "core/apr/apr.c": """## 概述
外观(Appearance)系统，集中管理所有游戏对象的视觉模板。

### apr_init
初始化15种APR模板(4玩家+6子弹+2敌人+1金币+1火墙)，设置尺寸/碰撞箱/图片名/透明度，然后循环调用load_img_dsc加载图片到内存

### apr_get
根据枚举ID查表返回模板指针

### apr_apply
将模板应用到游戏对象：更新apr指针、LVGL图像源、调试碰撞箱尺寸
""",

        "core/pool/pool.c": """## 概述
通用对象池实现。用外部提供的静态数组作为空闲索引栈，O(1)分配/回收。

### pool_init
设置容量，将所有索引压入空闲栈

### pool_alloc
弹出栈顶索引返回；栈空返回POOL_INVALID_ID

### pool_free
校验后压回索引(防重复归还)

### pool_free_count
返回stack_top(剩余空闲数)
""",

        "core/timer/timer.c": """## 概述
游戏定时器，基于对象池管理30个定时器，支持ONCE和REPEAT模式。

### timer_create
从池分配→填充owner/interval/callback→记录start_tick=play_tick_get()

### timer_update (每帧调用)
遍历所有活跃定时器→检查owner存活→计算elapsed→到期则调callback；ONCE模式回收，REPEAT重置start_tick
""",

        "core/input_sw/input_sw.c": """## 概述
按键输入分发层。将硬件按键扫描结果通过短按/长按/持续按下三种机制分发。

### input_dispatch
input_hw_scan()→input_sw_dispatch()

### 三种回调注册
- press_callback: 短按(按下→抬起触发)
- long_press_callback: 长按(按住150ms后循环触发)
- key_down_callback: 持续按下(按下后立即循环，用于射击和技能CD)

### input_sw_dispatch (static)
遍历所有按键→对pressed调用press_callbacks→对long_press调用long_press_callbacks(非阻塞延时控制频率)→对key_down调用key_down_callbacks
""",

        "level/level.c": """## 概述
关卡管理器，5关难度递增，每关多波次(普通+Boss)。

### level_init
初始化5关波次配置静态数组

### level_update (每帧)
管理波次状态机：WAVE_PENDING(等待wave_delay)→WAVE_SPAWNING(逐个生成敌人)→waiting_cleanup(等Boss死或等清理延迟)→推进下一波

### 波次类型
- WAVE_TYPE_NORMAL: 按间隔逐个生成，最后一个是大怪
- WAVE_TYPE_BOSS: 立即生成单个Boss(高HP+高伤害)
""",

        "save/save.c": """## 概述
存档系统，key=value格式，支持PC(stdio)和MCU(FatFS)双平台。

### save_load
try_parse_file(主文件)→失败则try_parse_file(.backup)→都失败则apply_defaults()→成功后写回

### save_write
备份旧存档→写入8项数据(coin_num, vol_bgm, vol_sfx, vol_amp, show_hitbox, selected_plane, unlocked_mask, draw_count)

### save_clear
删除存档文件→重置为默认值
""",
    }

    for fname, content in files.items():
        add_section(d, fname, content)

    d.save(f"{OUT}/Beat_Plane_游戏核心.docx")
    print("游戏核心.docx 已生成")

# ============================================================
if __name__ == "__main__":
    make_overview()
    make_core_docs()
    print("Done!")
